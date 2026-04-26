# app/parsers/structured_word_parser.py
import io
import asyncio
from typing import List, Dict, Any
from docx import Document as DocxDocument
from .base_parser import FileParserStrategy


class StructuredWordParser(FileParserStrategy):
    """
    结构化Word文档解析器
    
    核心功能:
    1. 样式分析 -> 识别标题样式
    2. 段落层级 -> 构建文档结构
    3. 表格提取 -> 结构化表格数据
    4. 列表处理 -> 保留列表格式
    """
    
    def get_supported_mime_types(self) -> List[str]:
        return [
            "application/msword",  # .doc
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"  # .docx
        ]
    
    async def parse(self, file_bytes: bytes) -> str:
        """
        解析Word文件，提取结构化内容
        
        Args:
            file_bytes: Word文件的字节流
            
        Returns:
            str: 结构化的Markdown格式文本
        """
        import sys
        print(f"[WordParser] 接收文件大小: {len(file_bytes)} bytes", file=sys.stderr)
        if not self.validate_file(file_bytes):
            print(f"[WordParser] 文件验证失败: file_bytes={len(file_bytes)}", file=sys.stderr)
            raise ValueError("Word文件为空或无效")
        
        structured_content, extracted_images = await asyncio.to_thread(
            self._extract_structured_content, file_bytes
        )
        
        if not structured_content.strip():
            raise ValueError("Word文件内容为空")
        
        if extracted_images:
            print(f"[WordParser] 检测到 {len(extracted_images)} 张图片，开始VLM识别...", file=sys.stderr)
            structured_content = await self._process_images_with_vlm(
                structured_content, extracted_images
            )
        
        return structured_content.strip()
    
    def _extract_structured_content(self, file_bytes: bytes) -> tuple:
        """同步结构化内容提取（在线程池中运行）"""
        import sys
        import zipfile
        try:
            file_stream = io.BytesIO(file_bytes)
            doc = DocxDocument(file_stream)
            
            self._file_bytes = file_bytes
            
            print(f"[WordParser] 文档段落总数: {len(doc.paragraphs)}", file=sys.stderr)
            print(f"[WordParser] 文档表格总数: {len(doc.tables)}", file=sys.stderr)
            
            style_hierarchy = self._analyze_style_hierarchy(doc)
            print(f"[WordParser] 检测到的样式层级: {style_hierarchy}", file=sys.stderr)
            
            structured_blocks, extracted_images = self._extract_structured_blocks(doc, style_hierarchy)
            
            markdown_content = self._build_markdown(structured_blocks)
            
            total_chars = len(markdown_content)
            total_words = len(markdown_content.split())
            print(f"[WordParser] 提取完成 - 总字符数: {total_chars}, 总词数: {total_words}", file=sys.stderr)
            print(f"[WordParser] 提取图片数: {len(extracted_images)}", file=sys.stderr)
            
            del self._file_bytes
            
            return markdown_content, extracted_images
            
        except (ValueError, KeyError) as e:
            raise Exception(f"结构化Word解析数据错误: {str(e)}")
        except (OSError, IOError) as e:
            raise Exception(f"结构化Word解析IO错误: {str(e)}")
        except Exception as e:
            raise Exception(f"结构化Word解析失败: {str(e)}")
    
    def _analyze_style_hierarchy(self, doc) -> Dict[str, int]:
        """
        分析Word样式层级，推断标题级别
        
        算法思路:
        1. 识别内置标题样式 (Heading 1, Heading 2, ...)
        2. 分析字体大小和样式特征
        3. 构建样式名称到标题级别的映射
        
        Returns:
            Dict[str, int]: 样式名称 -> 标题级别的映射
        """
        style_hierarchy = {}
        
        # 内置标题样式映射
        builtin_headings = {
            'Heading 1': 1, 'Heading 2': 2, 'Heading 3': 3,
            'Heading 4': 4, 'Heading 5': 5, 'Heading 6': 6,
            '标题 1': 1, '标题 2': 2, '标题 3': 3,
            '标题 4': 4, '标题 5': 5, '标题 6': 6
        }
        
        # 收集所有段落样式信息
        style_info = {}
        
        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name
            
            if style_name not in style_info:
                style_info[style_name] = {
                    'font_size': None,
                    'is_bold': False,
                    'count': 0,
                    'avg_length': 0
                }
            
            style_info[style_name]['count'] += 1
            
            # 分析字体特征
            if paragraph.runs:
                first_run = paragraph.runs[0]
                if first_run.font.size:
                    font_size = first_run.font.size.pt
                    if style_info[style_name]['font_size'] is None:
                        style_info[style_name]['font_size'] = font_size
                
                if first_run.font.bold:
                    style_info[style_name]['is_bold'] = True
            
            # 计算平均长度
            text_length = len(paragraph.text.strip())
            style_info[style_name]['avg_length'] = (
                style_info[style_name]['avg_length'] + text_length
            ) / 2
        
        # 1. 首先添加内置标题样式
        for style_name, level in builtin_headings.items():
            if style_name in style_info:
                style_hierarchy[style_name] = level
        
        # 2. 分析其他可能的标题样式
        # 标题特征：字体大、加粗、文本短、出现次数少
        for style_name, info in style_info.items():
            if style_name in style_hierarchy:
                continue
            
            # 启发式规则判断是否为标题
            is_heading = (
                info['is_bold'] and  # 加粗
                info['avg_length'] < 50 and  # 文本较短
                info['count'] < 20  # 出现次数不多
            )
            
            if is_heading:
                # 根据字体大小推断级别
                font_size = info.get('font_size', 12)
                if font_size >= 18:
                    level = 1
                elif font_size >= 16:
                    level = 2
                elif font_size >= 14:
                    level = 3
                else:
                    level = 4
                
                style_hierarchy[style_name] = level
        
        return style_hierarchy
    
    def _extract_structured_blocks(self, doc, style_hierarchy: Dict[str, int]) -> tuple:
        """
        提取结构化文档块
        
        Returns:
            tuple: (structured_blocks, extracted_images)
        """
        import sys
        
        structured_blocks = []
        extracted_images = []
        empty_paragraphs = 0
        paragraphs_with_images = 0
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            style_name = paragraph.style.name
            
            has_images = False
            image_index = 0
            for run in paragraph.runs:
                if run._element.xpath('.//w:drawing') or run._element.xpath('.//w:pict'):
                    has_images = True
                    break
                image_index += 1
            
            if not text and not has_images:
                empty_paragraphs += 1
                continue
            
            image_placeholder = ""
            image_bytes = None
            if has_images:
                try:
                    image_placeholder, image_bytes = self._extract_image_content_from_paragraph(paragraph)
                    if image_bytes:
                        placeholder_idx = len(extracted_images)
                        extracted_images.append({
                            "id": placeholder_idx,
                            "bytes": image_bytes,
                            "placeholder": f"[IMAGE OCR PLACEHOLDER_{placeholder_idx}]"
                        })
                        image_placeholder = f"\n\n{extracted_images[-1]['placeholder']}\n"
                        print(f"[WordParser] 检测到图片，已提取字节数据，大小: {len(image_bytes)} bytes", file=sys.stderr)
                except Exception as e:
                    print(f"[WordParser] 图片提取失败: {e}", file=sys.stderr)
                    image_placeholder = ""
                    image_bytes = None
            
            combined_content = text
            if image_placeholder:
                if combined_content:
                    combined_content += f"\n\n{image_placeholder}"
                else:
                    combined_content = image_placeholder
            
            if not combined_content.strip():
                paragraphs_with_images += 1
                continue
            
            if style_name in style_hierarchy:
                block = {
                    "type": "heading",
                    "level": style_hierarchy[style_name],
                    "content": combined_content
                }
                if image_bytes:
                    block["has_image"] = True
                structured_blocks.append(block)
            else:
                if self._is_list_paragraph(paragraph):
                    block = {
                        "type": "list_item",
                        "content": combined_content
                    }
                    if image_bytes:
                        block["has_image"] = True
                    structured_blocks.append(block)
                else:
                    block = {
                        "type": "paragraph",
                        "content": combined_content
                    }
                    if image_bytes:
                        block["has_image"] = True
                    structured_blocks.append(block)
        
        image_blocks, standalone_images = self._extract_standalone_images(doc)
        if image_blocks:
            structured_blocks.extend(image_blocks)
            extracted_images.extend(standalone_images)
            print(f"[WordParser] 提取独立图片块: {len(image_blocks)}", file=sys.stderr)
        
        for table in doc.tables:
            table_content = self._extract_table_content(table)
            if table_content:
                structured_blocks.append({
                    "type": "table",
                    "content": table_content
                })
        
        print(f"[WordParser] 空段落数: {empty_paragraphs}", file=sys.stderr)
        print(f"[WordParser] 只含图片段落数: {paragraphs_with_images}", file=sys.stderr)
        print(f"[WordParser] 段落总数: {len(doc.paragraphs)}", file=sys.stderr)
        print(f"[WordParser] 提取的块数: {len(structured_blocks)}", file=sys.stderr)
        print(f"[WordParser] 处理图片数: {len(extracted_images)}", file=sys.stderr)
        
        return structured_blocks, extracted_images
    
    def _extract_image_content_from_paragraph(self, paragraph) -> tuple:
        """
        从段落中提取图片内容
        
        Args:
            paragraph: Word段落对象
            
        Returns:
            tuple: (placeholder_text, image_bytes or None)
        """
        import sys
        import zipfile
        import os
        
        try:
            image_count = 0
            image_descriptions = []
            image_bytes = None
            
            for run in paragraph.runs:
                drawings = run._element.xpath('.//w:drawing')
                for drawing in drawings:
                    image_count += 1
                    
                    extent = drawing.find('.//a:ext', namespaces={
                        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
                    })
                    
                    size_info = ""
                    if extent is not None:
                        cx = extent.get('cx', '0')
                        cy = extent.get('cy', '0')
                        try:
                            width_inches = int(cx) / 914400
                            height_inches = int(cy) / 914400
                            size_info = f"({width_inches:.1f}\" x {height_inches:.1f}\")"
                        except (ValueError, TypeError):
                            size_info = ""
                    
                    alt_text = ""
                    cNvPr = drawing.find('.//p:cNvPr', namespaces={
                        'p': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
                    })
                    if cNvPr is not None:
                        alt_text = cNvPr.get('descr', '') or cNvPr.get('name', '')
                    
                    desc = f"[图片 {image_count}]"
                    if alt_text:
                        desc = f"[图片: {alt_text}]"
                    elif size_info:
                        desc = f"[图片 {size_info}]"
                    
                    image_descriptions.append(desc)
                    
                    if image_bytes is None:
                        image_bytes = self._extract_image_from_drawing(drawing)
                
                picts = run._element.xpath('.//w:pict')
                for pict in picts:
                    image_count += 1
                    image_descriptions.append(f"[图片 {image_count}]")
                    if image_bytes is None:
                        image_bytes = self._extract_image_from_pict(pict)
            
            if not image_descriptions:
                return "", None
            
            placeholder = "\n\n[IMAGE CONTENT]\n"
            for desc in image_descriptions:
                placeholder += f"- {desc}\n"
            
            return placeholder, image_bytes
            
        except Exception as e:
            print(f"[WordParser] 图片提取异常: {e}", file=sys.stderr)
            return "", None
    
    def _extract_image_from_drawing(self, drawing) -> bytes:
        """从w:drawing元素提取图片字节"""
        import sys
        
        try:
            blip = drawing.find('.//a:blip', namespaces={
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
            })
            
            if blip is not None:
                embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if embed_id:
                    return self._get_image_from_relationship(embed_id)
            
            inline = drawing.find('.//wp:inline', namespaces={
                'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
            })
            if inline is not None:
                blip = inline.find('.//a:blip', namespaces={
                    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
                })
                if blip is not None:
                    embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed_id:
                        return self._get_image_from_relationship(embed_id)
            
            anchor = drawing.find('.//wp:anchor', namespaces={
                'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
            })
            if anchor is not None:
                blip = anchor.find('.//a:blip', namespaces={
                    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
                })
                if blip is not None:
                    embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed_id:
                        return self._get_image_from_relationship(embed_id)
            
            return None
            
        except Exception as e:
            print(f"[WordParser] drawing图片提取失败: {e}", file=sys.stderr)
            return None
    
    def _extract_image_from_pict(self, pict) -> bytes:
        """从w:pict元素提取图片字节"""
        import sys
        
        try:
            shapetype = pict.find('.//v:shapetype')
            if shapetype is not None:
                pass
            
            image_elem = pict.find('.//v:image', namespaces={
                'v': 'urn:schemas-microsoft-com:vml'
            })
            if image_elem is not None:
                olink = image_elem.get('{urn:schemas-microsoft-com:office:office}link')
                if olink:
                    return self._get_image_from_ole(olink)
            
            return None
            
        except Exception as e:
            print(f"[WordParser] pict图片提取失败: {e}", file=sys.stderr)
            return None
    
    def _get_image_from_relationship(self, embed_id: str) -> bytes:
        """从文档关系中获取图片字节"""
        import sys
        import zipfile
        import io
        
        try:
            file_bytes = getattr(self, '_file_bytes', None)
            if file_bytes is None:
                return None
            
            with zipfile.ZipFile(io.BytesIO(file_bytes), 'r') as zf:
                doc_xml = zf.read('word/_rels/document.xml.rels')
                
                import re
                pattern = f'Id="{embed_id}"[^>]*Target="([^"]*)"'
                match = re.search(pattern, doc_xml.decode('utf-8'))
                
                if match:
                    image_path = 'word/' + match.group(1)
                    if image_path.startswith('word/'):
                        image_path = image_path[5:]
                    image_path = 'word/' + image_path
                    
                    try:
                        return zf.read(image_path)
                    except KeyError:
                        return None
                
                return None
            
        except Exception as e:
            print(f"[WordParser] 图片关系提取失败: {e}", file=sys.stderr)
            return None
    
    def _get_image_from_ole(self, olink: str) -> bytes:
        """从OLE链接获取图片"""
        return None
    
    def _get_doc_part(self):
        """获取文档part"""
        try:
            return self._doc_part
        except AttributeError:
            return None
    
    def _extract_standalone_images(self, doc) -> tuple:
        """
        提取文档中独立的图片（不在任何文本段落中的图片）
        
        Args:
            doc: Word文档对象
            
        Returns:
            tuple: (image_blocks, extracted_images)
        """
        import sys
        
        image_blocks = []
        extracted_images = []
        
        try:
            for para_idx, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                has_images = False
                image_bytes = None
                
                for run in para.runs:
                    if run._element.xpath('.//w:drawing') or run._element.xpath('.//w:pict'):
                        has_images = True
                        break
                
                if has_images and not text:
                    image_block = {
                        "type": "paragraph",
                        "content": f"[IMAGE: 文档中的嵌入图片 {para_idx + 1}]"
                    }
                    image_blocks.append(image_block)
                    
                    for run in para.runs:
                        drawings = run._element.xpath('.//w:drawing')
                        for drawing in drawings:
                            if image_bytes is None:
                                image_bytes = self._extract_image_from_drawing(drawing)
                                if image_bytes:
                                    placeholder_idx = len(extracted_images)
                                    extracted_images.append({
                                        "id": placeholder_idx,
                                        "bytes": image_bytes,
                                        "placeholder": f"[IMAGE OCR PLACEHOLDER_{placeholder_idx}]"
                                    })
                                    print(f"[WordParser] 检测到独立图片，字节数据，大小: {len(image_bytes)} bytes", file=sys.stderr)
                                    break
                        if image_bytes:
                            break
                    
                    print(f"[WordParser] 检测到独立图片块，段落索引: {para_idx + 1}", file=sys.stderr)
            
            return image_blocks, extracted_images
            
        except Exception as e:
            print(f"[WordParser] 独立图片提取异常: {e}", file=sys.stderr)
            return [], []
    
    async def _process_images_with_vlm(
        self, 
        markdown_content: str, 
        extracted_images: List[Dict[str, Any]]
    ) -> str:
        """
        使用VLM处理图片，生成描述并替换占位符
        
        Args:
            markdown_content: Markdown格式的文档内容
            extracted_images: 提取的图片列表
            
        Returns:
            str: 替换图片占位符后的内容
        """
        import sys
        
        if not extracted_images:
            return markdown_content
        
        try:
            from app.services.vlm_service import vlm_service
            
            if not vlm_service.is_enabled:
                print("[WordParser] VLM服务未启用，跳过图片OCR", file=sys.stderr)
                return markdown_content
            
            print(f"[WordParser] 开始VLM处理 {len(extracted_images)} 张图片...", file=sys.stderr)
            
            processed_content = markdown_content
            
            for img_info in extracted_images:
                placeholder = img_info.get("placeholder", "")
                image_bytes = img_info.get("bytes")
                
                if not placeholder or not image_bytes:
                    continue
                
                try:
                    print(f"[WordParser] 处理图片 {img_info['id'] + 1}/{len(extracted_images)}...", file=sys.stderr)
                    
                    image_description = await vlm_service.describe_image(image_bytes)
                    
                    replacement = f"\n\n[图片内容说明]:\n{image_description}\n\n"
                    
                    processed_content = processed_content.replace(placeholder, replacement)
                    
                    print(f"[WordParser] 图片 {img_info['id'] + 1} 处理完成，描述长度: {len(image_description)}", file=sys.stderr)
                    
                except Exception as e:
                    print(f"[WordParser] 图片 {img_info['id']} VLM处理失败: {e}", file=sys.stderr)
                    replacement = f"\n\n[图片内容 - VLM处理失败]\n\n"
                    processed_content = processed_content.replace(placeholder, replacement)
            
            print(f"[WordParser] VLM图片处理完成", file=sys.stderr)
            return processed_content
            
        except ImportError:
            print("[WordParser] VLM服务导入失败，跳过图片OCR", file=sys.stderr)
            return markdown_content
        except Exception as e:
            print(f"[WordParser] VLM处理异常: {e}", file=sys.stderr)
            return markdown_content
    
    def _is_list_paragraph(self, paragraph) -> bool:
        """判断段落是否为列表项"""
        text = paragraph.text.strip()
        
        # 检查是否以列表标记开头
        list_markers = ['•', '·', '-', '*', '○', '□', '■']
        for marker in list_markers:
            if text.startswith(marker):
                return True
        
        # 检查是否以数字编号开头
        if text and text[0].isdigit():
            # 查找数字后的标点
            for i, char in enumerate(text):
                if char in '.、)':
                    return True
                elif not char.isdigit():
                    break
        
        return False
    
    def _extract_table_content(self, table) -> str:
        """提取表格内容并格式化为Markdown"""
        try:
            table_data = []
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace('\n', ' ')
                    row_data.append(cell_text)
                table_data.append(row_data)
            
            if not table_data:
                return ""
            
            # 格式化为Markdown表格
            markdown_lines = []
            
            # 表头
            if table_data:
                header = table_data[0]
                markdown_lines.append("| " + " | ".join(header) + " |")
                markdown_lines.append("| " + " | ".join(["---"] * len(header)) + " |")
                
                # 表格内容
                for row in table_data[1:]:
                    markdown_lines.append("| " + " | ".join(row) + " |")
            
            return "\n".join(markdown_lines)
            
        except (ValueError, KeyError) as e:
            print(f"表格提取数据错误: {str(e)}")
            return ""
        except (OSError, IOError) as e:
            print(f"表格提取IO错误: {str(e)}")
            return ""
        except Exception as e:
            print(f"表格提取失败: {str(e)}")
            return ""
    
    def _build_markdown(self, structured_blocks: List[Dict[str, Any]]) -> str:
        """将结构化块转换为Markdown格式"""
        markdown_lines = []
        current_list_level = 0
        
        for block in structured_blocks:
            if block["type"] == "heading":
                # 标题：添加对应数量的#
                level = block["level"]
                heading_prefix = "#" * level
                markdown_lines.append(f"{heading_prefix} {block['content']}")
                markdown_lines.append("")  # 标题后空行
                current_list_level = 0
                
            elif block["type"] == "paragraph":
                # 段落：直接添加内容
                markdown_lines.append(block["content"])
                markdown_lines.append("")  # 段落后空行
                current_list_level = 0
                
            elif block["type"] == "list_item":
                # 列表项：添加Markdown列表格式
                content = block["content"]
                
                # 移除原有的列表标记
                for marker in ['•', '·', '-', '*', '○', '□', '■']:
                    if content.startswith(marker):
                        content = content[1:].strip()
                        break
                
                # 移除数字编号
                if content and content[0].isdigit():
                    for i, char in enumerate(content):
                        if char in '.、)':
                            content = content[i+1:].strip()
                            break
                
                markdown_lines.append(f"- {content}")
                current_list_level = 1
                
            elif block["type"] == "table":
                # 表格：添加表格内容
                if current_list_level > 0:
                    markdown_lines.append("")  # 列表后需要空行
                markdown_lines.append(block["content"])
                markdown_lines.append("")  # 表格后空行
                current_list_level = 0
        
        return "\n".join(markdown_lines)