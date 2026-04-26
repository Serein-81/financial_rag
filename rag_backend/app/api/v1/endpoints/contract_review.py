"""
合同审核智能助手 API 端点
提供合同深度分析和风险评估的 RESTful 接口
"""

import logging
from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from typing import Optional, List
from io import BytesIO
import uuid
from datetime import datetime

from app.api.deps import get_current_user, CurrentUser
from app.schemas.contract_review import (
    ContractAnalysisRequest,
    ContractAnalysisResponse,
    DeepClauseAnalysisRequest,
    DeepClauseAnalysisResponse,
    ContractComparisonRequest,
    ContractComparisonResponse,
    ContractType,
    ClauseType,
    RiskLevel,
)
from app.services.contract_review_service import ContractReviewService
from app.services.pdf_export_service import pdf_export_service
from app.services.file_service import file_service
from app.services.minio_service import minio_service

router = APIRouter(prefix="/contract-review", tags=["合同审核"])
logger = logging.getLogger(__name__)

contract_review_service = ContractReviewService()


@router.post("/analyze", response_model=ContractAnalysisResponse)
async def analyze_contract(
    request: ContractAnalysisRequest,
    user: CurrentUser = Depends(get_current_user),
):
    """
    分析合同
    
    对合同文本进行全文分析，包括：
    - 条款提取和分类
    - 风险评估
    - 关键发现
    - 修改建议
    """
    try:
        request.user_id = str(user.id)
        request.tenant_id = user.tenant_id
        
        result = await contract_review_service.analyze_contract(request)
        
        return ContractAnalysisResponse(**result)
        
    except (ValueError, KeyError) as e:
        logger.error(f"❌ 合同分析数据错误: {e}", exc_info=True)
        raise HTTPException(status_code=400, detail=f"分析数据错误: {str(e)}")
    except (OSError, IOError) as e:
        logger.error(f"❌ 合同分析IO错误: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"分析IO错误: {str(e)}")
    except (OSError, IOError) as e:
        raise HTTPException(status_code=500, detail=f"IO错误: {str(e)}")
    except Exception as e:
        logger.error(f"❌ 合同分析失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"分析失败: {str(e)}")


@router.post("/upload")
async def upload_and_analyze_contract(
    file: UploadFile = File(..., description="合同文件（支持 PDF、Word、文本等格式）"),
    contract_name: str = Form(..., description="合同名称"),
    contract_type: ContractType = Form(ContractType.OTHER, description="合同类型"),
    counterparty_name: Optional[str] = Form(None, description="对方名称"),
    contract_value: Optional[float] = Form(None, description="合同金额"),
    effective_date: Optional[str] = Form(None, description="生效日期"),
    expiration_date: Optional[str] = Form(None, description="到期日期"),
    user: CurrentUser = Depends(get_current_user),
):
    """
    上传合同并智能分析
    
    支持文件上传，自动解析合同内容，调用法务智能体进行深度分析
    并生成专业的风险评估报告。
    
    **功能特点**：
    - 支持多种文件格式（PDF、Word、文本等）
    - 自动提取合同文本内容
    - 调用法务智能体进行专业分析
    - 生成风险评估报告
    - 提供修改建议
    
    **使用场景**：
    - 合同签订前审核
    - 合同条款风险评估
    - 合同合规性检查
    - 合同谈判支持
    """
    try:
        # 1. 验证文件类型
        allowed_types = [
            "application/pdf",
            "application/msword",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "text/plain",
            "text/markdown",
            "application/octet-stream"
        ]
        
        if file.content_type not in allowed_types:
            raise HTTPException(
                status_code=400,
                detail=f"不支持的文件类型: {file.content_type}。"
                      f"支持的格式: PDF、Word (.doc, .docx)、文本文件"
            )
        
        # 2. 读取文件内容
        file_content = await file.read()
        logger.info(f"📄 接收到合同文件: {file.filename}, 大小: {len(file_content)} bytes")
        
        # 3. 生成文件存储路径
        file_id = str(uuid.uuid4())
        file_extension = file.filename.split('.')[-1] if '.' in file.filename else 'bin'
        minio_path = f"contracts/{user.tenant_id}/{file_id}.{file_extension}"
        
        # 4. 上传到 MinIO（使用正确的参数）
        await minio_service.upload_document_async(
            file_bytes=file_content,
            object_name=minio_path,
            content_type=file.content_type,
            tenant_id=user.tenant_id
        )
        logger.info(f"✅ 文件已上传到 MinIO: {minio_path}")
        
        # 5. 提取合同文本内容（参考税务提交流程）
        try:
            # 先尝试直接解析文件内容
            contract_text = await _extract_document_text(file_content, file.content_type)
            
            if not contract_text or len(contract_text.strip()) < 50:
                logger.warning("直接解析失败，尝试使用 OCR 服务...")
                contract_text = await _extract_with_ocr_service(file_content, file.content_type)
            
            if not contract_text or len(contract_text.strip()) < 50:
                raise HTTPException(
                    status_code=422,
                    detail="无法从文件中提取文本内容，请确保文件包含可识别的文本。"
                )
            
            logger.info(f"📝 成功提取合同文本，长度: {len(contract_text)} 字符")
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"❌ 文件解析失败: {e}", exc_info=True)
            raise HTTPException(
                status_code=422,
                detail=f"文件解析失败: {str(e)}。请确保文件格式正确且内容可读。"
            )
        
        # 6. 构建分析请求
        from datetime import date as date_type
        analysis_request = ContractAnalysisRequest(
            tenant_id=user.tenant_id,
            user_id=str(user.id),
            contract_name=contract_name,
            contract_type=contract_type,
            contract_content=contract_text,
            counterparty_name=counterparty_name,
            contract_value=contract_value,
            effective_date=date_type.fromisoformat(effective_date) if effective_date else None,
            expiration_date=date_type.fromisoformat(expiration_date) if expiration_date else None,
            include_deep_analysis=True,
            include_risk_assessment=True,
            include_suggestions=True
        )
        
        # 7. 调用法务智能体进行深度分析
        logger.info(f"🔍 正在调用法务智能体分析合同: {contract_name}")
        result = await contract_review_service.analyze_contract_with_legal_agent(
            request=analysis_request,
            user_id=str(user.id),
            tenant_id=user.tenant_id
        )
        
        # 8. 保存文件元数据
        file_metadata = {
            "file_id": file_id,
            "file_name": file.filename,
            "minio_path": minio_path,
            "content_type": file.content_type,
            "size": len(file_content),
            "uploaded_by": str(user.id),
            "uploaded_at": datetime.utcnow().isoformat(),
            "analysis_id": result.get("analysis_id")
        }
        
        logger.info(f"✅ 合同分析完成，analysis_id: {result.get('analysis_id')}")
        
        return {
            "success": True,
            "message": "合同上传并分析成功",
            "analysis_id": result.get("analysis_id"),
            "file_metadata": file_metadata,
            "result": result
        }
        
    except HTTPException:
        raise
    except (ValueError, KeyError) as e:
        logger.error(f"❌ 合同上传数据错误: {e}", exc_info=True)
        raise HTTPException(status_code=400, detail=f"上传数据错误: {str(e)}")
    except (OSError, IOError) as e:
        logger.error(f"❌ 合同上传IO错误: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"上传IO错误: {str(e)}")
    except Exception as e:
        logger.error(f"❌ 合同上传分析失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"上传分析失败: {str(e)}")


async def _extract_document_text(content: bytes, content_type: str) -> str:
    """
    从文档中提取文本（参考税务提交流程）
    
    Args:
        content: 文件内容
        content_type: 文件类型
    
    Returns:
        提取的文本内容
    """
    try:
        # PDF 文件
        if 'pdf' in content_type.lower() or content[:4] == b'%PDF':
            return await _extract_pdf_text(content)
        
        # Word 文件
        elif 'wordprocessingml' in content_type.lower() or 'msword' in content_type.lower():
            return await _extract_word_text(content)
        
        # 文本文件
        elif 'text/plain' in content_type.lower() or 'text/markdown' in content_type.lower():
            return content.decode('utf-8', errors='ignore')
        
        else:
            # 尝试作为文本解析
            try:
                return content.decode('utf-8', errors='ignore')
            except:
                return ""
                
    except Exception as e:
        logger.error(f"❌ 文档文本提取失败: {e}", exc_info=True)
        return ""


async def _extract_pdf_text(content: bytes) -> str:
    """提取 PDF 文本层"""
    try:
        try:
            from pypdf import PdfReader
            logger.info("使用 pypdf 库提取 PDF 文本")
        except ImportError:
            import PyPDF2
            PdfReader = PyPDF2.PdfReader
            logger.info("使用 PyPDF2 库提取 PDF 文本")
        
        import io
        pdf_file = io.BytesIO(content)
        pdf_reader = PdfReader(pdf_file)
        text_parts = []
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_parts.append(page_text)
        return "\n".join(text_parts)
    except ImportError:
        logger.warning("pypdf/PyPDF2 未安装，尝试 pdfplumber")
        try:
            import pdfplumber
            import io
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                text_parts = [page.extract_text() for page in pdf.pages if page.extract_text()]
            return "\n".join(text_parts)
        except ImportError:
            logger.error("无法提取 PDF 内容：缺少 pypdf/PyPDF2 或 pdfplumber")
            return ""
        except Exception as e:
            logger.error(f"PDF 文本提取失败: {str(e)}")
            return ""
    except Exception as e:
        logger.error(f"PDF 文本提取失败: {str(e)}")
        return ""


async def _extract_word_text(content: bytes) -> str:
    """提取 Word 文档文本"""
    try:
        import io
        from docx import Document
        
        doc_file = io.BytesIO(content)
        doc = Document(doc_file)
        
        # 提取段落文本
        paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        
        # 提取表格文本
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        tables_text.append(cell.text.strip())
        
        full_text = "\n".join(paragraphs)
        if tables_text:
            full_text += "\n\n[表格内容]\n" + "\n".join(tables_text)
        
        return full_text
        
    except ImportError:
        logger.error("python-docx 未安装，无法解析 Word 文档")
        return ""
    except Exception as e:
        logger.error(f"Word 文档文本提取失败: {str(e)}")
        return ""


async def _extract_with_ocr_service(content: bytes, content_type: str) -> str:
    """使用 OCR 服务提取文本（扫描件/图片）"""
    try:
        from app.services.ocr_factory import OCRFactory
        
        ocr_factory = OCRFactory()
        available_engines = ocr_factory.available_engines
        
        logger.info(f"[OCR] 开始 OCR 处理，可用引擎: {available_engines}")
        
        if not available_engines:
            logger.warning("[OCR] 没有可用的 OCR 引擎")
            return ""
        
        # 检测是否为 PDF 文件
        is_pdf = content[:4] == b'%PDF'
        
        if is_pdf:
            # PDF 文件：先尝试 Unstructured API
            if 'unstructured' in available_engines:
                logger.info("[OCR] 尝试使用 Unstructured API")
                adapter = ocr_factory.get_adapter('unstructured')
                
                if adapter and hasattr(adapter, 'extract_text'):
                    import tempfile
                    import os
                    
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                        tmp.write(content)
                        tmp_path = tmp.name
                    
                    try:
                        result = await adapter.extract_text(tmp_path)
                        if result and len(result.strip()) > 50:
                            logger.info(f"[OCR] Unstructured API 提取成功: {len(result)} 字符")
                            return result
                    finally:
                        try:
                            os.unlink(tmp_path)
                        except:
                            pass
            
            # 如果 Unstructured API 失败，尝试 PDF 转图片 OCR
            logger.info("[OCR] 尝试将 PDF 转换为图片进行 OCR")
            try:
                import fitz
                
                pdf_doc = fitz.open(stream=content, filetype="pdf")
                total_pages = len(pdf_doc)
                
                all_text = []
                
                for page_num in range(total_pages):
                    page = pdf_doc[page_num]
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    img_bytes = pix.tobytes("png")
                    
                    # 对每一页进行 OCR
                    page_text = await _ocr_image(img_bytes)
                    if page_text:
                        all_text.append(page_text)
                
                pdf_doc.close()
                
                result = "\n\n".join(all_text)
                logger.info(f"[OCR] PDF OCR 完成，提取 {len(result)} 字符")
                return result
                
            except ImportError:
                logger.warning("PyMuPDF 未安装")
                return ""
        else:
            # 图片文件直接 OCR
            return await _ocr_image(content)
        
    except Exception as e:
        logger.error(f"[OCR] OCR 服务处理失败: {e}", exc_info=True)
        return ""


async def _ocr_image(image_bytes: bytes) -> str:
    """对图片进行 OCR 识别"""
    try:
        from app.services.ocr_factory import OCRFactory
        
        ocr_factory = OCRFactory()
        available_engines = ocr_factory.available_engines
        
        if not available_engines:
            return ""
        
        # 优先使用 PaddleOCR
        if 'paddleocr' in available_engines:
            adapter = ocr_factory.get_adapter('paddleocr')
            if adapter and hasattr(adapter, 'extract_text'):
                return await adapter.extract_text(image_bytes)
        
        # 其次使用 Tesseract
        if 'tesseract' in available_engines:
            adapter = ocr_factory.get_adapter('tesseract')
            if adapter and hasattr(adapter, 'extract_text'):
                return await adapter.extract_text(image_bytes)
        
        return ""
        
    except Exception as e:
        logger.error(f"[OCR] 图片 OCR 失败: {e}")
        return ""


@router.post("/clause-analysis", response_model=DeepClauseAnalysisResponse)
async def analyze_clause_deeply(
    request: DeepClauseAnalysisRequest,
    user: CurrentUser = Depends(get_current_user),
):
    """
    深度条款分析
    
    对单个条款进行深度分析，包括：
    - 法律解释
    - 潜在问题识别
    - 行业惯例对比
    - 修改建议
    - 相关法规参考
    """
    try:
        request.user_id = str(user.id)
        request.tenant_id = user.tenant_id
        
        result = await contract_review_service.analyze_clause_deeply(request)
        
        return DeepClauseAnalysisResponse(**result)
        
    except (ValueError, KeyError) as e:
        logger.error(f"❌ 深度条款分析数据错误: {e}", exc_info=True)
        raise HTTPException(status_code=400, detail=f"分析数据错误: {str(e)}")
    except (OSError, IOError) as e:
        logger.error(f"❌ 深度条款分析IO错误: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"分析IO错误: {str(e)}")
    except (OSError, IOError) as e:
        raise HTTPException(status_code=500, detail=f"IO错误: {str(e)}")
    except Exception as e:
        logger.error(f"❌ 深度条款分析失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"分析失败: {str(e)}")


@router.post("/compare", response_model=ContractComparisonResponse)
async def compare_contracts(
    request: ContractComparisonRequest,
    user: CurrentUser = Depends(get_current_user),
):
    """
    对比合同
    
    对比两个合同的条款差异，包括：
    - 条款对比
    - 关键差异识别
    - 优势分析
    - 风险对比
    - 谈判要点建议
    """
    try:
        request.user_id = str(user.id)
        request.tenant_id = user.tenant_id
        
        result = await contract_review_service.compare_contracts(request)
        
        return ContractComparisonResponse(**result)
        
    except (ValueError, KeyError) as e:
        logger.error(f"❌ 合同对比数据错误: {e}", exc_info=True)
        raise HTTPException(status_code=400, detail=f"对比数据错误: {str(e)}")
    except (OSError, IOError) as e:
        logger.error(f"❌ 合同对比IO错误: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"对比IO错误: {str(e)}")
    except (OSError, IOError) as e:
        raise HTTPException(status_code=500, detail=f"IO错误: {str(e)}")
    except Exception as e:
        logger.error(f"❌ 合同对比失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"对比失败: {str(e)}")


@router.get("/risks")
async def get_risk_assessments(
    analysis_id: str = Query(..., description="分析ID"),
    risk_level: Optional[str] = Query(None, description="风险级别筛选"),
    user: CurrentUser = Depends(get_current_user),
):
    """
    获取风险评估结果
    
    获取指定合同分析的风险评估详情
    """
    try:
        if analysis_id not in contract_review_service._analysis_cache:
            raise HTTPException(status_code=404, detail="分析结果不存在")
        
        analysis = contract_review_service._analysis_cache[analysis_id]
        
        risk_assessments = analysis.get("risk_assessments", [])
        
        if risk_level:
            try:
                risk_level_enum = RiskLevel(risk_level)
                risk_assessments = [
                    r for r in risk_assessments
                    if r.get("risk_level") == risk_level_enum.value
                ]
            except ValueError:
                raise HTTPException(status_code=400, detail="无效的风险级别")
        
        return {
            "analysis_id": analysis_id,
            "risk_assessments": risk_assessments,
            "total_count": len(risk_assessments),
            "high_risk_count": sum(
                1 for r in risk_assessments
                if r.get("risk_level") in [RiskLevel.HIGH.value, RiskLevel.CRITICAL.value]
            )
        }
        
    except HTTPException:
        raise
    except (ValueError, KeyError) as e:
        logger.error(f"❌ 获取风险评估数据错误: {e}", exc_info=True)
        raise HTTPException(status_code=400, detail=f"获取数据错误: {str(e)}")
    except (OSError, IOError) as e:
        logger.error(f"❌ 获取风险评估IO错误: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"获取IO错误: {str(e)}")
    except (OSError, IOError) as e:
        raise HTTPException(status_code=500, detail=f"IO错误: {str(e)}")
    except Exception as e:
        logger.error(f"❌ 获取风险评估失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"获取失败: {str(e)}")


@router.get("/clauses")
async def get_contract_clauses(
    analysis_id: str = Query(..., description="分析ID"),
    clause_type: Optional[str] = Query(None, description="条款类型筛选"),
    risk_level: Optional[str] = Query(None, description="风险级别筛选"),
    user: CurrentUser = Depends(get_current_user),
):
    """
    获取合同条款列表
    
    获取指定合同分析的条款详情
    """
    try:
        if analysis_id not in contract_review_service._analysis_cache:
            raise HTTPException(status_code=404, detail="分析结果不存在")
        
        analysis = contract_review_service._analysis_cache[analysis_id]
        clauses = analysis.get("clauses_extracted", [])
        
        if clause_type:
            try:
                clause_type_enum = ClauseType(clause_type)
                clauses = [
                    c for c in clauses
                    if c.get("clause_type") == clause_type_enum.value
                ]
            except ValueError:
                raise HTTPException(status_code=400, detail="无效的条款类型")
        
        if risk_level:
            try:
                risk_level_enum = RiskLevel(risk_level)
                clauses = [
                    c for c in clauses
                    if c.get("risk_level") == risk_level_enum.value
                ]
            except ValueError:
                raise HTTPException(status_code=400, detail="无效的风险级别")
        
        return {
            "analysis_id": analysis_id,
            "clauses": clauses,
            "total_count": len(clauses),
            "high_risk_count": sum(
                1 for c in clauses
                if c.get("risk_level") in [RiskLevel.HIGH.value, RiskLevel.CRITICAL.value]
            )
        }
        
    except HTTPException:
        raise
    except (ValueError, KeyError) as e:
        logger.error(f"❌ 获取条款列表数据错误: {e}", exc_info=True)
        raise HTTPException(status_code=400, detail=f"获取数据错误: {str(e)}")
    except (OSError, IOError) as e:
        logger.error(f"❌ 获取条款列表IO错误: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"获取IO错误: {str(e)}")
    except (OSError, IOError) as e:
        raise HTTPException(status_code=500, detail=f"IO错误: {str(e)}")
    except Exception as e:
        logger.error(f"❌ 获取条款列表失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"获取失败: {str(e)}")


@router.get("/clause-types")
async def list_clause_types():
    """
    获取条款类型列表
    
    返回所有支持的条款类型
    """
    return {
        "clause_types": [
            {"value": ct.value, "label": _get_clause_type_label(ct)}
            for ct in ClauseType
        ]
    }


@router.get("/report/export")
async def export_contract_review_report_pdf(
    analysis_id: str = Query(..., description="分析ID"),
    user: CurrentUser = Depends(get_current_user),
):
    """
    导出合同审核报告为 PDF
    
    根据分析ID导出合同审核报告 PDF
    
    Args:
        analysis_id: 分析ID
        
    Returns:
        PDF 文件的流式响应
    """
    try:
        if analysis_id not in contract_review_service._analysis_cache:
            raise HTTPException(status_code=404, detail="分析结果不存在")
        
        analysis = contract_review_service._analysis_cache[analysis_id]
        
        pdf_bytes = pdf_export_service.export_contract_review_report(analysis)
        
        filename = f"contract_review_report_{analysis_id}.pdf"
        
        return StreamingResponse(
            BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{filename}"
            }
        )
        
    except HTTPException:
        raise
    except (ValueError, KeyError) as e:
        logger.error(f"❌ PDF导出数据错误: {e}", exc_info=True)
        raise HTTPException(status_code=400, detail=f"PDF导出数据错误: {str(e)}")
    except (OSError, IOError) as e:
        logger.error(f"❌ PDF导出IO错误: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"PDF导出IO错误: {str(e)}")
    except (OSError, IOError) as e:
        raise HTTPException(status_code=500, detail=f"IO错误: {str(e)}")
    except Exception as e:
        logger.error(f"❌ PDF导出失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"PDF导出失败: {str(e)}")


def _get_clause_type_label(clause_type: ClauseType) -> str:
    """获取条款类型标签"""
    labels = {
        ClauseType.PAYMENT: "付款条款",
        ClauseType.DELIVERY: "交付条款",
        ClauseType.WARRANTY: "保修条款",
        ClauseType.LIABILITY: "责任条款",
        ClauseType.TERMINATION: "终止条款",
        ClauseType.CONFIDENTIALITY: "保密条款",
        ClauseType.INTELLECTUAL_PROPERTY: "知识产权条款",
        ClauseType.DISPUTE_RESOLUTION: "争议解决条款",
        ClauseType.FORCE_MAJEURE: "不可抗力条款",
        ClauseType.INDEMNIFICATION: "赔偿条款",
        ClauseType.ASSIGNMENT: "转让条款",
        ClauseType.GOVERNING_LAW: "适用法律条款",
        ClauseType.OTHER: "其他条款",
    }
    return labels.get(clause_type, clause_type.value)


@router.get("/contract-types")
async def list_contract_types():
    """
    获取合同类型列表
    
    返回所有支持的合同类型
    """
    return {
        "contract_types": [
            {"value": ct.value, "label": _get_contract_type_label(ct)}
            for ct in ContractType
        ]
    }


def _get_contract_type_label(contract_type: ContractType) -> str:
    """获取合同类型标签"""
    labels = {
        ContractType.SALES: "销售合同",
        ContractType.PURCHASE: "采购合同",
        ContractType.SERVICE: "服务合同",
        ContractType.LABOR: "劳动合同",
        ContractType.LEASE: "租赁合同",
        ContractType.LOAN: "借款合同",
        ContractType.PARTNERSHIP: "合作协议",
        ContractType.CONFIDENTIALITY: "保密协议",
        ContractType.OTHER: "其他合同",
    }
    return labels.get(contract_type, contract_type.value)


@router.get("/templates")
async def list_contract_templates(
    contract_type: Optional[str] = Query(None, description="合同类型过滤"),
    user: CurrentUser = Depends(get_current_user),
):
    """
    获取合同审核模板列表
    
    返回预定义的合同审核模板，支持按类型过滤
    """
    templates = [
        {
            "id": "template-001",
            "name": "标准采购合同模板",
            "description": "适用于企业间标准采购业务，包含完整条款和风险提示",
            "contract_type": "purchase",
            "usage_count": 156
        },
        {
            "id": "template-002",
            "name": "销售合同标准模板",
            "description": "适用于产品销售场景，包含交付、质量保证等核心条款",
            "contract_type": "sales",
            "usage_count": 203
        },
        {
            "id": "template-003",
            "name": "服务外包合同模板",
            "description": "适用于专业服务外包，包含服务范围、验收标准等条款",
            "contract_type": "service",
            "usage_count": 178
        },
        {
            "id": "template-004",
            "name": "房屋租赁合同模板",
            "description": "适用于办公或商业租赁，包含租金、押金、装修等条款",
            "contract_type": "lease",
            "usage_count": 89
        },
        {
            "id": "template-005",
            "name": "劳动合同标准模板",
            "description": "适用于企业招聘员工，包含薪酬、福利、保密等条款",
            "contract_type": "employment",
            "usage_count": 312
        },
        {
            "id": "template-006",
            "name": "合作协议模板",
            "description": "适用于战略合作或项目合作，包含权益分配等条款",
            "contract_type": "partnership",
            "usage_count": 67
        },
        {
            "id": "template-007",
            "name": "借款合同模板",
            "description": "适用于企业或个人借贷，包含利率、还款方式等条款",
            "contract_type": "loan",
            "usage_count": 45
        },
        {
            "id": "template-008",
            "name": "保密协议模板",
            "description": "适用于商业机密保护，包含保密范围、违约责任等条款",
            "contract_type": "confidentiality",
            "usage_count": 234
        },
        {
            "id": "template-009",
            "name": "软件开发合同模板",
            "description": "适用于软件定制开发，包含需求变更、知识产权等条款",
            "contract_type": "service",
            "usage_count": 98
        },
        {
            "id": "template-010",
            "name": "供应链采购合同模板",
            "description": "适用于供应链采购，包含交货、质量控制等条款",
            "contract_type": "purchase",
            "usage_count": 76
        },
        {
            "id": "template-011",
            "name": "咨询服务合同模板",
            "description": "适用于各类咨询业务，包含服务内容、成果交付等条款",
            "contract_type": "service",
            "usage_count": 112
        },
        {
            "id": "template-012",
            "name": "设备租赁合同模板",
            "description": "适用于设备租赁业务，包含租金、维护、损坏赔偿等条款",
            "contract_type": "lease",
            "usage_count": 34
        }
    ]
    
    if contract_type:
        templates = [t for t in templates if t["contract_type"] == contract_type]
    
    return {
        "templates": templates,
        "total": len(templates)
    }


@router.get("/risk-levels")
async def list_risk_levels():
    """
    获取风险级别列表
    
    返回所有风险级别定义
    """
    return {
        "risk_levels": [
            {"value": rl.value, "label": _get_risk_level_label(rl), "description": _get_risk_level_description(rl)}
            for rl in RiskLevel
        ]
    }


def _get_risk_level_label(risk_level: RiskLevel) -> str:
    """获取风险级别标签"""
    labels = {
        RiskLevel.LOW: "低风险",
        RiskLevel.MEDIUM: "中等风险",
        RiskLevel.HIGH: "高风险",
        RiskLevel.CRITICAL: "极高风险",
    }
    return labels.get(risk_level, risk_level.value)


def _get_risk_level_description(risk_level: RiskLevel) -> str:
    """获取风险级别描述"""
    descriptions = {
        RiskLevel.LOW: "基本无风险，可正常执行",
        RiskLevel.MEDIUM: "存在一定风险，需要注意",
        RiskLevel.HIGH: "存在较高风险，建议谨慎处理",
        RiskLevel.CRITICAL: "存在严重风险，需要立即处理",
    }
    return descriptions.get(risk_level, "")


@router.get("/history")
async def get_analysis_history(
    page: int = Query(1, ge=1),
    page_size: int = Query(10, ge=1, le=100),
    contract_type: Optional[str] = Query(None),
    risk_level: Optional[str] = Query(None),
    user: CurrentUser = Depends(get_current_user),
):
    """
    获取合同审核历史记录
    """
    try:
        service = ContractReviewService()
        result = await service.get_analysis_history(
            user_id=str(user.id),
            tenant_id=user.tenant_id,
            page=page,
            page_size=page_size,
            contract_type=contract_type,
            risk_level=risk_level
        )
        return result
    except Exception as e:
        logger.error(f"❌ 获取合同审核历史失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"获取合同审核历史失败: {str(e)}")


@router.get("/health")
async def health_check():
    """健康检查端点"""
    return {
        "status": "healthy",
        "service": "contract_review",
        "version": "1.0.0"
    }
